"""报告导出：Markdown → HTML（纯 Python）/ DOCX（python-docx）"""
import html
import re

def md_to_html(md: str) -> str:
    """简易 Markdown → HTML：标题/表格/列表/引用/代码块/行内格式。转义防 XSS"""
    lines = md.split("\n")
    out = []
    i = 0
    in_table = False
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # 代码块
        if stripped.startswith("```"):
            if in_code:
                out.append("<pre><code>" + html.escape("\n".join(code_buf)) + "</code></pre>")
                code_buf, in_code = [], False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        # 表格
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|?$", lines[i + 1].strip()):
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            out.append("<table><thead><tr>" + "".join(f"<th>{_inline(h)}</th>" for h in headers) + "</tr></thead><tbody>")
            i += 2
            in_table = True
            continue
        if in_table:
            if stripped.startswith("|"):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                out.append("<tr>" + "".join(f"<td>{_inline(c)}</td>" for c in cells) + "</tr>")
            else:
                out.append("</tbody></table>")
                in_table = False
                continue
            i += 1
            continue
        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            lvl = len(m.group(1))
            out.append(f"<h{lvl}>{_inline(m.group(2))}</h{lvl}>")
            i += 1
            continue
        # 引用
        if stripped.startswith(">"):
            out.append(f"<blockquote>{_inline(stripped[1:].strip())}</blockquote>")
            i += 1
            continue
        # 列表
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            items = [m.group(1)]
            i += 1
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.match(r"^[-*]\s+", lines[i].strip()).group(0) and re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            out.append("<ul>" + "".join(f"<li>{_inline(it)}</li>" for it in items) + "</ul>")
            continue
        # 分隔线/空行
        if re.match(r"^-{3,}$", stripped) or not stripped:
            out.append("")
            i += 1
            continue
        out.append(f"<p>{_inline(stripped)}</p>" if stripped else "")
        i += 1
    if in_table:
        out.append("</tbody></table>")
    if in_code:
        out.append("<pre><code>" + html.escape("\n".join(code_buf)) + "</code></pre>")
    return "\n".join(out)

def _inline(text: str) -> str:
    """行内格式：代码/加粗/链接 + HTML 转义"""
    text = html.escape(text)
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r'<a href="\2">\1</a>', text)
    return text

def export_docx(md: str, out_path: str) -> None:
    """Markdown → DOCX（python-docx；需 pip install python-docx）"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise SystemExit("未安装 python-docx，请执行: pip install python-docx")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    in_code = False
    code_buf = []
    for line in md.split("\n"):
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_buf))
                p.style = doc.styles["Normal"]
                code_buf, in_code = [], False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if stripped.startswith("|") and not stripped.replace("|", "").replace("-", "").replace(":", "").strip():
            continue  # 分隔行
        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if len(doc.tables) and len(doc.tables[-1].rows) and len(doc.tables[-1].rows[-1].cells) == len(cells):
                doc.tables[-1].add_row()
                row = doc.tables[-1].rows[-1]
                for j, c in enumerate(cells):
                    row.cells[j].text = re.sub(r"[*`]", "", c)
            else:
                t = doc.add_table(rows=1, cols=len(cells))
                for j, c in enumerate(cells):
                    t.rows[0].cells[j].text = re.sub(r"[*`]", "", c)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            doc.add_heading(re.sub(r"[*`]", "", m.group(2)), level=min(len(m.group(1)), 4))
            continue
        if stripped.startswith(">"):
            doc.add_paragraph(re.sub(r"[*`]", "", stripped[1:].strip()))
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "• ", re.sub(r"[*`]", "", stripped)))
            continue
        if stripped:
            doc.add_paragraph(re.sub(r"[*`]", "", stripped))
    doc.save(out_path)
