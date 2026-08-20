import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))
import pytest
from export import md_to_html, export_docx

SAMPLE_MD = """# 标题一

> 引用内容

| 概念 | 解释 |
|------|------|
| A | 甲 |

- 列表项1
- 列表项2

**加粗** 和 `行内代码`
"""

def test_md_to_html_basic():
    html = md_to_html(SAMPLE_MD)
    assert "<h1>标题一</h1>" in html
    assert "<blockquote>引用内容</blockquote>" in html
    assert "<table>" in html and "甲" in html
    assert "<ul>" in html and "列表项1" in html
    assert "<strong>加粗</strong>" in html
    assert "<code>行内代码</code>" in html

def test_md_to_html_escape():
    html = md_to_html("# <script>alert(1)</script>")
    assert "<script>" not in html  # XSS 转义
    assert "&lt;script&gt;" in html

def test_export_docx(tmp_path):
    out = tmp_path / "r.docx"
    export_docx(SAMPLE_MD, str(out))
    assert out.exists()
    from docx import Document
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "标题一" in texts
    assert "甲" in " ".join(texts) or any("甲" in c.text for t in doc.tables for r in t.rows for c in r.cells)
